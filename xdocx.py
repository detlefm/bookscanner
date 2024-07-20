from docx import Document
from docx.shared import Pt
import os
from pathlib import Path
import shutil
from xfile import getfiles
#from constants import DOC_TEMPLATE
from config import config


def docx_from_template(filename:str, append:bool=False):
    if os.path.exists(filename):
        if append:
            return Document(filename)
        else:
            os.remove(filename)
    templatename = config.getvalue('doc_template')
    if os.path.exists(templatename):
        shutil.copy(templatename,filename)
        return Document(filename)
    return Document()



def text_to_word( word_file:str,txtlst:list[str], pagebreak:bool=True, append:bool=False):
    # Ggfls. erstellen eines neuen Word-Dokuments
    doc = docx_from_template(word_file,append=append)
    for txt in txtlst:
        doc.add_paragraph(text=txt)
        if pagebreak:
            doc.add_page_break()
    # Speichern des Word-Dokuments
    doc.save(word_file)



if __name__ == "__main__":
    import sys
    if len(sys.argv)<3:
        print(f'usage: {sys.argv[0]}  pathtoimages outfilename')
        exit(-1)
    files = getfiles(sys.argv[1],suffixes=['.txt'])
    txtlst = [f.read_text(encoding='utf-8') for f in files]
    text_to_word(word_file=sys.argv[2],txtlst=txtlst)
    